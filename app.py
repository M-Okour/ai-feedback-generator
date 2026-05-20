import io
import re
import time
import zipfile
import pandas as pd
import streamlit as st

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openai import OpenAI, RateLimitError, APIError, APITimeoutError


st.set_page_config(page_title="AI Feedback Generator", layout="wide")
st.title("AI Student Feedback Generator from Marks and Rubric")
assessor_name = st.text_input("Assessor Name")
feedback_mode = st.radio(
    "Feedback generation mode",
    options=[
        "Generate feedback for all PCs",
        "Generate feedback only for Not Yet Competent PCs"
    ],
    index=0
)
sa2_date = st.text_input(
    "Second Attempt Date",
    placeholder="Example: 15/06/2026"
)

signature_file = st.file_uploader(
    "Upload Assessor Signature Image",
    type=["png", "jpg", "jpeg"]
)

signature_date = st.text_input(
    "Signature Date",
    placeholder="Example: 20/05/2026"
)

signature_bytes = None

if signature_file is not None:
    signature_file.seek(0)
    signature_bytes = signature_file.read()

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def extract_student_signatures_from_excel(classlist_file):
    """
    Extracts embedded signature images from the Excel column named Student_Signature.
    Returns:
    {
        "H00123456": image_bytes
    }
    """

    classlist_file.seek(0)
    wb = load_workbook(classlist_file)
    ws = wb.active

    header_row = 1
    student_id_col = None
    signature_col = None

    for cell in ws[header_row]:
        header = str(cell.value).strip() if cell.value else ""

        if "id" in header.lower():
            student_id_col = cell.column

        if header == "Student_Signature":
            signature_col = cell.column

    if student_id_col is None or signature_col is None:
        return {}

    signatures = {}

    for image in ws._images:
        anchor = image.anchor._from
        image_col = anchor.col + 1
        image_row = anchor.row + 1

        if image_col == signature_col:
            student_id = ws.cell(row=image_row, column=student_id_col).value

            if student_id:
                student_id = str(student_id).strip()
                signatures[student_id] = image._data()

    return signatures

def insert_image_in_cell(cell, image_bytes, width_inches=1.1):
    if not image_bytes:
        return

    cell.text = ""
    run = cell.paragraphs[0].add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))

def fill_student_signature_fields(table, student_signature_bytes, signature_date):
    if not student_signature_bytes and not signature_date:
        return

    signature_done = False
    student_signature_done = False

    for row_index, row in enumerate(table.rows):
        cells = row.cells

        for i, cell in enumerate(cells):
            text = cell.text.strip().lower()

            # First place: adjacent to "Signature:"
            if (
                not signature_done
                and text in ["signature:", "signature"]
            ):
                if student_signature_bytes and i + 1 < len(cells):
                    insert_image_in_cell(cells[i + 1], student_signature_bytes)

                if signature_date:
                    for j, date_cell in enumerate(cells):
                        if "date" in date_cell.text.strip().lower():
                            if j + 1 < len(cells):
                                cells[j + 1].text = signature_date
                            break

                signature_done = True

            # Second place: adjacent to "Student Signature:"
            if "student signature" in text and not student_signature_done:
                if student_signature_bytes:
                    fill_adjacent_or_empty(
                        row=row,
                        label_keywords=["Student Signature"],
                        value=""
                    )

                    for j in range(i + 1, len(cells)):
                        if is_empty_cell(cells[j]):
                            insert_image_in_cell(cells[j], student_signature_bytes)
                            break

                # following row date
                if signature_date and row_index + 1 < len(table.rows):
                    n
                    
                    ext_row = table.rows[row_index + 1]
                    fill_adjacent_or_empty(
                        row=next_row,
                        label_keywords=["Date"],
                        value=signature_date
                    )

                student_signature_done = True
def insert_signature_image(cell, signature_bytes, width_inches=1.2):
    if not signature_bytes:
        return

    cell.text = ""

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(signature_bytes), width=Inches(width_inches))
    

def fill_signature_fields(table, signature_bytes, signature_date):
    if not signature_bytes and not signature_date:
        return

    second_signature_done = False

    for row_index, row in enumerate(table.rows):
        cells = row.cells

        for i, cell in enumerate(cells):
            text = cell.text.strip().lower()

            # -------------------------------------------------
            # First location:
            # Assessor Signature:
            # Use fill_adjacent_or_empty for date/text fields
            # -------------------------------------------------
            if "assessor signature" in text:

                if signature_bytes:
                    # Put signature image in adjacent/empty cell
                    for j in range(i + 1, len(cells)):
                        if is_empty_cell(cells[j]):
                            insert_signature_image(cells[j], signature_bytes)
                            break
                    else:
                        if i + 1 < len(cells):
                            insert_signature_image(cells[i + 1], signature_bytes)

                if signature_date and row_index + 1 < len(table.rows):
                    next_row = table.rows[row_index + 1]

                    fill_adjacent_or_empty(
                        row=next_row,
                        label_keywords=["Date"],
                        value=signature_date
                    )

            # -------------------------------------------------
            # Second location ONLY ONCE:
            # Keep original logic for Signature: and Date
            # -------------------------------------------------
            elif (
                not second_signature_done
                and text in ["signature:", "signature"]
            ):

                if signature_bytes and i + 1 < len(cells):
                    insert_signature_image(cells[i + 1], signature_bytes)

                if signature_date:
                    for j, date_cell in enumerate(cells):
                        if "date" in date_cell.text.strip().lower():
                            if j + 1 < len(cells):
                                cells[j + 1].text = signature_date
                            break

                second_signature_done = True


def extract_lo_sections_from_rubric(rubric_text):
    """
    Extracts LO/Element sections from rubric.

    Returns:
    {
        "LO1": {
            "title": "Use vector algebra to solve force and moment problems",
            "pcs": ["PC1.1", "PC1.2", "PC1.3"]
        },
        "LO2": {
            "title": "Analyse equilibrium problems",
            "pcs": ["PC2.1", "PC2.2"]
        }
    }
    """

    lo_data = {}

    element_pattern = r"Element\s+(\d+)\s*:\s*(.*?)(?=Element\s+\d+\s*:|$)"
    element_matches = re.finditer(element_pattern, rubric_text, flags=re.IGNORECASE | re.DOTALL)

    for match in element_matches:
        element_number = match.group(1)
        section = match.group(0)
        title_line = match.group(2).strip().split("\n")[0].strip()

        pcs = []

        for pc_match in re.finditer(r"\bPC\s*(\d+\.\d+)\b", section, flags=re.IGNORECASE):
            pc = f"PC{pc_match.group(1)}"
            if pc not in pcs:
                pcs.append(pc)

        lo_data[f"LO{element_number}"] = {
            "title": title_line,
            "pcs": pcs
        }

    return lo_data

def build_lo_comments(feedback_rows, lo_data):
    """
    Builds general LO comments based on PC marks under each Element/LO.
    """

    pc_lookup = {
        normalize_pc_for_matching(row["PC"]): row
        for row in feedback_rows
    }

    comments = []

    for lo, data in lo_data.items():
        pcs = data["pcs"]
        title = data["title"]

        related_rows = [
            pc_lookup[normalize_pc_for_matching(pc)]
            for pc in pcs
            if normalize_pc_for_matching(pc) in pc_lookup
        ]

        if not related_rows:
            continue

        failed = [
            row["PC"]
            for row in related_rows
            if row["Level"] == "Not Yet Competent"
        ]

        levels = [row["Level"] for row in related_rows]

        if failed:
            comment = (
                f"{lo}: In {title}, further improvement is required in "
                f"{', '.join(failed)}. Please review the relevant methods, show clear working steps, "
                f"and check the final answers carefully."
            )
        elif all(level == "Competent with Distinction" for level in levels):
            comment = (
                f"{lo}: In {title}, you demonstrated strong achievement across the related performance criteria, "
                f"with clear methods, accurate solutions, and well-presented work."
            )
        elif any(level == "Competent with Merit" for level in levels):
            comment = (
                f"{lo}: In {title}, you demonstrated good achievement across the related performance criteria. "
                f"To improve further, focus on completing all details such as labels, units, and final verification."
            )
        else:
            comment = (
                f"{lo}: In {title}, you demonstrated a competent level of understanding across the related performance criteria. "
                f"Further improvement can be made by improving accuracy, presentation, and completeness of solution steps."
            )

        comments.append(comment)

    return comments

        
def fill_assessor_name_in_table(table, assessor_name):
    """
    Replaces:
    Assessor Name:

    with:
    Assessor Name: John Smith
    """

    if not assessor_name:
        return

    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()

            if "Assessor Name:" in text:
                cell.text = f"Assessor Name: {assessor_name}"
        
def get_level(mark):
    mark = float(mark)
    if mark < 60:
        return "Not Yet Competent"
    elif mark < 70:
        return "Competent"
    elif mark < 85:
        return "Competent with Merit"
    return "Competent with Distinction"

def get_first_name(full_name):
    if pd.isna(full_name):
        return "Student"
    full_name = str(full_name).strip()
    return full_name.split()[0] if full_name else "Student"


def normalize_pc_for_matching(text):
    text = str(text).upper().replace(" ", "")

    match = re.search(r"E(\d+):?PC(\d+)$", text)
    if match:
        return f"PC{match.group(1)}.{match.group(2)}"

    match = re.search(r"E\d+:?PC(\d+\.\d+)", text)
    if match:
        return f"PC{match.group(1)}"

    match = re.search(r"PC(\d+\.\d+)", text)
    if match:
        return f"PC{match.group(1)}"

    match = re.fullmatch(r"(\d+\.\d+)", text)
    if match:
        return f"PC{match.group(1)}"

    return text


def read_docx_text(file):
    file.seek(0)
    doc = Document(file)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_pc_list_from_rubric(rubric_text):
    pcs = []

    patterns = [
        r"\bE\s*(\d+)\s*:?\s*PC\s*(\d+\.\d+|\d+)\b",
        r"\bPC\s*(\d+\.\d+|\d+)\b",
        r"\b(\d+\.\d+)\s+(?:Determine|Analyse|Analyze|Calculate|Use|Construct|Apply)\b",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, rubric_text, flags=re.IGNORECASE):
            if len(match.groups()) == 2:
                element = match.group(1)
                pc_number = match.group(2)
                pc_code = f"PC{pc_number}" if "." in pc_number else f"PC{element}.{pc_number}"
            else:
                pc_number = match.group(1)
                pc_code = f"PC{pc_number}"

            pc_code = normalize_pc_for_matching(pc_code)

            if pc_code not in pcs:
                pcs.append(pc_code)

    return pcs


def extract_rubric_sections(rubric_text, pc_list):
    sections = {}
    headings = []

    for pc in pc_list:
        pc_num = pc.replace("PC", "")
        patterns = [
            rf"\bPC\s*{re.escape(pc_num)}\b",
            rf"\b{re.escape(pc_num)}\b",
        ]

        for pattern in patterns:
            for match in re.finditer(pattern, rubric_text, flags=re.IGNORECASE):
                headings.append({"pc": pc, "start": match.start()})

    headings = sorted(headings, key=lambda x: x["start"])

    clean = []
    for h in headings:
        if not clean:
            clean.append(h)
        else:
            last = clean[-1]
            if h["pc"] == last["pc"] and abs(h["start"] - last["start"]) < 50:
                continue
            clean.append(h)

    for i, h in enumerate(clean):
        start = h["start"]
        end = clean[i + 1]["start"] if i + 1 < len(clean) else len(rubric_text)
        sections[h["pc"]] = rubric_text[start:end].strip()

    for pc in pc_list:
        sections.setdefault(pc, "")

    return sections


def find_student_columns(df):
    name_col = None
    id_col = None

    for col in df.columns:
        c = str(col).lower()

        if "name" in c and name_col is None:
            name_col = col

        if "id" in c and id_col is None:
            id_col = col

    return name_col, id_col


def get_pc_columns(df):
    pc_cols = {}

    for col in df.columns:
        col_text = str(col).upper().replace(" ", "")
        normalized = normalize_pc_for_matching(col_text)

        if re.fullmatch(r"PC\d+\.\d+", normalized):
            pc_cols[normalized] = col

    return pc_cols

def get_overall_level_from_marks(pc_marks):
    marks = [float(m) for m in pc_marks.values()]

    if not marks:
        return "Not Yet Competent"

    if any(m < 60 for m in marks):
        return "Not Yet Competent"

    average = round(sum(marks) / len(marks))

    return get_level(average)


def extract_sa_number_from_doc(doc):
    """
    Looks for titles like:
    Summative Assessment 1
    SUMMATIVE ASSESSMENT 2
    SA1
    SA 2
    """

    full_text = []

    for p in doc.paragraphs:
        full_text.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    text = "\n".join(full_text)

    match = re.search(r"Summative\s+Assessment\s+(\d+)", text, flags=re.IGNORECASE)
    if match:
        return match.group(1)

    match = re.search(r"\bSA\s*(\d+)\b", text, flags=re.IGNORECASE)
    if match:
        return match.group(1)

    return ""

def build_summative_comment(
    student_name,
    overall_level,
    sa_number,
    feedback_rows,
    sa2_date,
    lo_data
):
    failed_pcs = [
        row["PC"]
        for row in feedback_rows
        if row["Level"] == "Not Yet Competent"
    ]

    first_name = get_first_name(student_name)

    # -------------------------------------------------
    # Main overall comment
    # -------------------------------------------------

    if failed_pcs:

        failed_text = ", ".join(failed_pcs)

        summary_comment = (
            f"{first_name}, you have achieved an overall "
            f"{overall_level} grade in SA{sa_number}. "
            f"Please see the feedback for {failed_text}."
        )

        if sa2_date.strip():
            summary_comment += (
                f" The second attempt will be conducted on "
                f"{sa2_date}."
            )

    else:

        summary_comment = (
            f"{first_name}, you have achieved an overall "
            f"{overall_level} grade in SA{sa_number}."
        )

    # -------------------------------------------------
    # LO comments
    # -------------------------------------------------

    lo_comments = build_lo_comments(
        feedback_rows=feedback_rows,
        lo_data=lo_data
    )

    return {
        "lo_comments": lo_comments,
        "summary_comment": summary_comment
    }

def generate_ai_feedback(first_name, pc, level, rubric_section, max_retries=4):
    prompt = f"""
You are an experienced engineering assessor writing feedback for a diploma student.

Student first name:
{first_name}

Performance Criterion:
{pc}

Student competency level:
{level}

Relevant rubric criteria:
{rubric_section}

Write feedback for the Assessor Feedback section.

Style requirements:
- Sound natural, human, and supportive.
- Keep a professional assessor tone.
- Use the student's first name once at the beginning.
- Write 1 to 2 sentences only.
- Do not sound robotic or generic.
- Mention one specific strength based on the level.
- Mention one clear improvement area.
- Do not mention AI, automated marking, rubric file, or the exact numerical mark.
- Do not use bullet points.
- Do not overpraise if the level is low.
- Keep the language suitable for official assessment feedback.
"""

    for attempt in range(max_retries):
        try:
            response = client.responses.create(
                model="gpt-4.1-mini",
                input=prompt,
                max_output_tokens=140
            )
            return response.output_text.strip()

        except RateLimitError:
            wait_time = 8 * (attempt + 1)
            st.warning(f"Rate limit reached while generating feedback for {pc}. Retrying in {wait_time} seconds...")
            time.sleep(wait_time)

        except (APIError, APITimeoutError):
            wait_time = 5 * (attempt + 1)
            st.warning(f"Temporary API issue for {pc}. Retrying in {wait_time} seconds...")
            time.sleep(wait_time)

        except Exception:
            break

    return (
        f"{first_name}, you achieved {level} for {pc}. "
        f"Please review the relevant solution steps, accuracy, presentation, and final answer to improve your performance."
    )


# =========================================================
# Template filling helpers - safer for merged cells
# =========================================================

def row_texts(row):
    return [cell.text.strip() for cell in row.cells]


def row_full_text(row):
    return " | ".join(row_texts(row))


def is_empty_cell(cell):
    return cell.text.strip() == ""


def fill_first_empty_cell_in_row(row, value, start_index=0):
    cells = row.cells

    for i in range(start_index, len(cells)):
        if is_empty_cell(cells[i]):
            cells[i].text = str(value)
            return True

    return False


def fill_adjacent_or_empty(row, label_keywords, value):
    """
    Safer than cells[i+1] for merged-cell templates.
    It tries:
    1. cell immediately after the label
    2. first empty cell after the label
    3. first empty cell in the row
    """

    cells = row.cells

    for i, cell in enumerate(cells):
        text = cell.text.strip()

        if any(keyword.lower() in text.lower() for keyword in label_keywords):

            # Try immediate right cell
            if i + 1 < len(cells) and is_empty_cell(cells[i + 1]):
                cells[i + 1].text = str(value)
                return True

            # Try any empty cell after label
            if fill_first_empty_cell_in_row(row, value, start_index=i + 1):
                return True

            # Try any empty cell in row
            if fill_first_empty_cell_in_row(row, value, start_index=0):
                return True

    return False


def fill_name_and_id_in_table(table, student_name, student_id):
    """
    Handles both:
    Student Name | [empty]
    ID No.       | [empty]

    Safer for merged tables because it searches empty target cells.
    """

    for row in table.rows:
        fill_adjacent_or_empty(
            row=row,
            label_keywords=["Student Name"],
            value=student_name
        )

        fill_adjacent_or_empty(
            row=row,
            label_keywords=["ID No.", "ID No", "Student ID"],
            value=student_id
        )


def get_row_pc(row, pc_marks):
    """
    Find whether this row belongs to one of the PCs.
    """

    for cell in row.cells:
        normalized = normalize_pc_for_matching(cell.text)

        if normalized in pc_marks:
            return normalized

    return None


def find_grade_column_indices(table):
    """
    Detect competency columns ONLY from the
    'PC Grade %' row.
    """

    for row in table.rows:

        cells = row.cells

        row_text = " | ".join(
            cell.text.strip().lower()
            for cell in cells
        )

        if "pc grade" not in row_text:
            continue

        grade_columns = {}

        for i, cell in enumerate(cells):

            text = (
                cell.text.strip()
                .lower()
                .replace("–", "-")
                .replace("—", "-")
            )

            if "0" in text and "59" in text:
                grade_columns["Not Yet Competent"] = i

            elif "60" in text and "69" in text:
                grade_columns["Competent"] = i

            elif "70" in text and "84" in text:
                grade_columns["Competent with Merit"] = i

            elif "85" in text and "100" in text:
                grade_columns["Competent with Distinction"] = i

        if grade_columns:
            return grade_columns



def fill_marks_in_assessment_table(table, pc_marks):
    """
    Search row by PC number and column by grade level,
    then fill the mark at the row-column intersection.
    """

    grade_columns = find_grade_column_indices(table)

    for row in table.rows:
        cells = row.cells

        # Find PC in this row
        row_pc = None

        for cell in cells:
            normalized_pc = normalize_pc_for_matching(cell.text)

            if normalized_pc in pc_marks:
                row_pc = normalized_pc
                break

        if row_pc is None:
            continue

        mark = pc_marks[row_pc]
        level = get_level(mark)

        if level not in grade_columns:
            continue

        target_col = grade_columns[level]

        if target_col < len(cells):
            # Optional: clear grade-band cells first
            for col in grade_columns.values():
                if col < len(cells):
                    cells[col].text = ""

            cells[target_col].text = str(int(mark))


def fill_summative_grade_in_table(table, pc_marks):
    marks = [float(m) for m in pc_marks.values()]

    if not marks:
        return

    if any(m < 60 for m in marks):
        summative = min(marks)
    else:
        summative = round(sum(marks) / len(marks))

    for row in table.rows:
        text = row_full_text(row)

        if "Summative Assessment Grade %:" in text:
            filled = fill_adjacent_or_empty(
                row=row,
                label_keywords=["Summative Assessment Grade %:"],
                value=int(summative)
            )

            return

def build_feedback_table_in_cell(
    cell,
    feedback_rows,
    student_name,
    overall_level,
    sa_number,
    sa2_date,
    lo_data
):
    cell.text = ""

    # -------------------------------------------------
    # Individual PC feedback
    # -------------------------------------------------

    for row_data in feedback_rows:

        pc = row_data["PC"]
        level = row_data["Level"]
        feedback = row_data["Feedback"]

        # Skip empty feedback rows if desired
        if not str(feedback).strip():
            continue

        # ---------------------------------------------
        # Heading
        # ---------------------------------------------

        heading = cell.add_paragraph()

        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1

        run = heading.add_run(f"{pc}: {level}:")
        run.bold = True
        run.font.size = Pt(9)

        # ---------------------------------------------
        # Feedback paragraph
        # ---------------------------------------------

        p = cell.add_paragraph(feedback)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1

        for r in p.runs:
            r.font.size = Pt(9)

    # -------------------------------------------------
    # Summative comment
    # -------------------------------------------------
    
    comment_data = build_summative_comment(
        student_name=student_name,
        overall_level=overall_level,
        sa_number=sa_number,
        feedback_rows=feedback_rows,
        sa2_date=sa2_date,
        lo_data=lo_data
    )
    
    # -------------------------------------------------
    # LO comments with bold LO1:, LO2:
    # -------------------------------------------------
    
    for lo_comment in comment_data["lo_comments"]:
    
        p = cell.add_paragraph()
    
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
    
        match = re.match(r"(LO\d+:)(.*)", lo_comment)
    
        if match:
    
            lo_label = match.group(1)
            lo_text = match.group(2).strip()
    
            run1 = p.add_run(lo_label + " ")
            run1.bold = True
            run1.font.size = Pt(9)
    
            run2 = p.add_run(lo_text)
            run2.font.size = Pt(9)
    
        else:
    
            run = p.add_run(lo_comment)
            run.font.size = Pt(9)
    
    # -------------------------------------------------
    # Final overall comment
    # -------------------------------------------------
    
    blank_p = cell.add_paragraph("")
    blank_p.paragraph_format.space_before = Pt(0)
    blank_p.paragraph_format.space_after = Pt(0)
    blank_p.paragraph_format.line_spacing = 1
    summary_p = cell.add_paragraph(
        comment_data["summary_comment"]
    )
    
    summary_p.paragraph_format.space_before = Pt(0)
    summary_p.paragraph_format.space_after = Pt(0)
    summary_p.paragraph_format.line_spacing = 1
    
    for r in summary_p.runs:
        r.font.size = Pt(10)
    
    return


def insert_feedback_table_at_assessor_feedback(
    doc,
    feedback_rows,
    student_name,
    overall_level,
    sa_number,
    sa2_date,
    lo_data
    ):
    feedback_inserted = False

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells

            for i, cell in enumerate(cells):
                if "Assessor Feedback:" in cell.text or "Assessor Feedback" in cell.text:
                    target_row_index = row._tr.getparent().index(row._tr) - 1

                    if target_row_index < len(table.rows):
                        next_row = table.rows[target_row_index]
                        # Prefer first empty cell in next row
                        target_cell = None
                
                        for next_cell in next_row.cells:
                            if is_empty_cell(next_cell):
                                target_cell = next_cell
                                break
                
                        # fallback = first cell of next row
                        if target_cell is None:
                            target_cell = next_row.cells[0]
                
                    else:
                        # fallback = current cell
                        target_cell = cell
                
                    build_feedback_table_in_cell(
                        target_cell,
                        feedback_rows,
                        student_name,
                        overall_level,
                        sa_number,
                        sa2_date,
                        lo_data
                    )
                
                    feedback_inserted = True
                    break

            if feedback_inserted:
                break

        if feedback_inserted:
            break


    if not feedback_inserted:

        p = doc.add_paragraph("Assessor Feedback:")
        p.runs[0].bold = True
    
        build_feedback_table_in_cell(
            cell=doc.add_table(rows=1, cols=1).cell(0, 0),
            feedback_rows=feedback_rows,
            student_name=student_name,
            overall_level=overall_level,
            sa_number=sa_number,
            sa2_date=sa2_date,
            lo_data=lo_data
        )

    return doc


def fill_template(doc,
                student_name,
                student_id,
                assessor_name,
                signature_bytes,
                student_signature_bytes,
                signature_date,
                feedback_rows,
                lo_data,
                sa2_date):
    pc_marks = {
        normalize_pc_for_matching(row["PC"]): row["Mark"]
        for row in feedback_rows
    }

    for table in doc.tables:
        table_text = "\n".join(
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
        )

        is_ack_table = (
            "Portfolio Evidence Requirements" in table_text
            or "Student Acknowledgment" in table_text
            or "Unit Title/s" in table_text
        )

        is_assessment_table = (
            "Assessment Results" in table_text
            or "PC Grade" in table_text
            or "Summative Assessment Grade" in table_text
            or "Grade Classification" in table_text
        )

        if is_ack_table:
            fill_name_and_id_in_table(table, student_name, student_id)
            fill_assessor_name_in_table(table, assessor_name)
            fill_signature_fields(table, signature_bytes, signature_date)
            fill_student_signature_fields(table, student_signature_bytes, signature_date)
            continue

        if is_assessment_table:
            fill_name_and_id_in_table(table, student_name, student_id)
            fill_marks_in_assessment_table(table, pc_marks)
            fill_summative_grade_in_table(table, pc_marks)
            fill_assessor_name_in_table(table, assessor_name)
            fill_signature_fields(table, signature_bytes, signature_date)
            fill_student_signature_fields(table, student_signature_bytes, signature_date)

    overall_level = get_overall_level_from_marks(pc_marks)
    sa_number = extract_sa_number_from_doc(doc)
                    
    
    doc = insert_feedback_table_at_assessor_feedback(
        doc=doc,
        feedback_rows=feedback_rows,
        student_name=student_name,
        overall_level=overall_level,
        sa_number=sa_number,
        sa2_date=sa2_date,
        lo_data=lo_data
    )

    fill_signature_fields(table, signature_bytes, signature_date)
        
    return doc


rubric_file = st.file_uploader("Upload rubric.docx", type=["docx"])
classlist_file = st.file_uploader("Upload classlist.xlsx", type=["xlsx"])
template_file = st.file_uploader("Upload Template_Feedback.docx", type=["docx"])
student_signatures = extract_student_signatures_from_excel(classlist_file)

if st.button("Generate AI Feedback Files"):
    if not rubric_file or not classlist_file or not template_file:
        st.error("Please upload rubric, classlist, and feedback template.")
        st.stop()

    rubric_text = read_docx_text(rubric_file)
    pc_list = extract_pc_list_from_rubric(rubric_text)
    rubric_sections = extract_rubric_sections(rubric_text, pc_list)
    
    lo_data = extract_lo_sections_from_rubric(rubric_text)
    
    st.write("LO/Element mapping:", lo_data)
    
    df = pd.read_excel(classlist_file)
    
    name_col, id_col = find_student_columns(df)
    pc_cols = get_pc_columns(df)
    
    st.subheader("Detected Setup")
    st.write("Name column:", name_col)
    st.write("ID column:", id_col)
    st.write("PCs from rubric:", pc_list)
    st.write("PC mark columns from classlist:", pc_cols)

    if not name_col or not id_col:
        st.error("Could not detect student name or ID column.")
        st.stop()

    if not pc_cols:
        st.error("Could not detect PC mark columns. Use headers like E1:PC1, PC1.1, PC3.1, or 3.1.")
        st.stop()

    zip_buffer = io.BytesIO()
    summary_rows = []
    progress = st.progress(0)

    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for index, student in df.iterrows():
            student_name = student[name_col]
            student_id = student[id_col]
            first_name = get_first_name(student_name)

            feedback_rows = []

            for pc in pc_list:
                normalized_pc = normalize_pc_for_matching(pc)

                if normalized_pc not in pc_cols:
                    continue

                mark = student[pc_cols[normalized_pc]]

                if pd.isna(mark):
                    continue

                mark = int(round(float(mark)))
                level = get_level(mark)
                rubric_section = rubric_sections.get(pc, "")

                should_generate_feedback = (
                    feedback_mode == "Generate feedback for all PCs"
                    or level == "Not Yet Competent"
                )

                if should_generate_feedback:
                    feedback = generate_ai_feedback(
                        first_name=first_name,
                        pc=pc,
                        level=level,
                        rubric_section=rubric_section
                    )
                else:
                    feedback = ""

                feedback_rows.append({
                    "PC": pc,
                    "Mark": mark,
                    "Level": level,
                    "Feedback": feedback
                })

                summary_rows.append({
                    "Student Name": student_name,
                    "Student ID": student_id,
                    "PC": pc,
                    "Mark": mark,
                    "Level": level,
                    "Feedback": feedback
                })

            template_file.seek(0)
            
            student_signature_bytes = student_signatures.get(str(student_id).strip())
            
            doc = Document(template_file)
            
            doc = fill_template(doc=doc,
                            student_name=student_name,
                            student_id=student_id,
                            assessor_name=assessor_name,
                            signature_bytes=signature_bytes,
                            student_signature_bytes=student_signature_bytes,
                            signature_date=signature_date,
                            feedback_rows=feedback_rows,
                            lo_data=lo_data,
                            sa2_date=sa2_date)

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            safe_name = re.sub(r'[\\/*?:"<>|]', "", str(student_name))
            safe_name = safe_name.replace(" ", "_")
            
            output_name = f"{student_id}_{safe_name}.docx"
            zip_file.writestr(output_name, doc_buffer.getvalue())

            progress.progress((index + 1) / len(df))

    zip_buffer.seek(0)

    st.success("AI feedback files generated successfully.")

    st.download_button(
        label="Download Feedback ZIP",
        data=zip_buffer.getvalue(),
        file_name="student_ai_feedback_files.zip",
        mime="application/zip"
    )

    if summary_rows:
        summary_df = pd.DataFrame(summary_rows)

        st.subheader("Feedback Summary")
        st.dataframe(summary_df)

        csv_buffer = io.StringIO()
        summary_df.to_csv(csv_buffer, index=False)

        st.download_button(
            label="Download Feedback Summary CSV",
            data=csv_buffer.getvalue(),
            file_name="feedback_summary.csv",
            mime="text/csv"
        )
